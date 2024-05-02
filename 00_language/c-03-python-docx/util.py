from copy import deepcopy
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.enum.text import WD_ALIGN_PARAGRAPH

import testdata

def docx_fill_data(wdoc):

  data = testdata.get_test_data()
  items = data['items']
  total = data['total']

  table = wdoc.tables[1]

  int_format = '{:,}'
  decimal_format = '{:,.2f}'
  for item in items:
    cells = table.add_row().cells
    cells[0].text = int_format.format(item['quantity'])
    cells[1].text = str(item['description'])
    cells[2].text = decimal_format.format(item['unitprice'])
    cells[3].text = decimal_format.format(item['linetotal'])

    for x in range(4):
      if (x == 1):
        continue

      cells[x].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT


  cells = table.add_row().cells
  cells[2].text = str('Total')
  cells[3].text = decimal_format.format(total)

  for x in range(2, 4):
    cell = cells[x]
    cell.paragraphs[0].runs[0].font.bold = True
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

  background = parse_xml(r'<w:shd {} w:fill="E7EEEE"/>'.format(nsdecls('w')))
  index = 0
  for row in table.rows:
    index += 1
    if (index == 1 or index % 2 == 0):
      continue

    for cell in row.cells:
      cell._tc.get_or_add_tcPr().append(deepcopy(background))