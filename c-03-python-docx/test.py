# https://python-docx.readthedocs.io/en/latest/user/quickstart.html

import os
import glob
from copy import deepcopy
from docx import Document
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.enum.text import WD_ALIGN_PARAGRAPH

from pprint import pprint

def clearDirectory(directory):
  files = glob.glob(f'{directory}*')
  for f in files:
    os.remove(f)

def test():
  template_file = './template/invoice-template.docx'
  result_directory = './result/'
  result_file = f'{result_directory}result.docx'

  clearDirectory(result_directory)

  wdoc = Document(template_file)
  table = wdoc.tables[1]
  
  items = [
    {
      'quantity': 12,
      'description': 'Item description No.0',
      'unitprice': 12000.3,
      'linetotal': 20
    }
  ]

  for x in range(10):
    items.append({
      'quantity': 1290,
      'description': f'Item description No.{x}',
      'unitprice': 12.3,
      'linetotal': 20000
    })

  # E7EEEE
  int_format = '{:,.0f}'
  decimal_format = '{:,.2f}'
  for item in items:
    cells = table.add_row().cells
    cells[0].text = int_format.format(item['quantity'])
    cells[1].text = str(item['description'])
    cells[2].text = decimal_format.format(item['unitprice'])
    cells[3].text = decimal_format.format(item['linetotal'])

    for x in range(4):
      if (x == 1):
        continue

      cells[x].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT


  cells = table.add_row().cells
  cells[0].text = ''
  cells[1].text = ''
  cells[2].text = str('Total')
  cells[3].text = decimal_format.format(50000.00)

  for x in range(4):
    cell = cells[x]
    cell.paragraphs[0].runs[0].font.bold = True
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT


  background = parse_xml(r'<w:shd {} w:fill="E7EEEE"/>'.format(nsdecls('w')))
  index = 0
  for row in table.rows:
    index += 1
    if (index == 1 or index % 2 == 0):
      continue

    for cell in row.cells:
      cell._tc.get_or_add_tcPr().append(deepcopy(background))

  wdoc.save(result_file)

  print('Done')

if __name__ == '__main__':
  test()